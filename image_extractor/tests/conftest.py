"""Shared fixtures.

Everything the suite reads is *built at runtime* rather than committed, so there are no
binary fixtures in the repository and the documents under test are known down to the pixel.

`temp_store` redirects the content-addressed store at a tmp path, so no test ever writes to
the real `store/` directory beside the source.
"""
import io
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import store  # noqa: E402


@pytest.fixture
def temp_store(tmp_path, monkeypatch):
    monkeypatch.setattr(store, "IMAGE_DIR", tmp_path / "store")
    return store


@pytest.fixture
def png_bytes():
    """A small real PNG. Built rather than committed so the suite has no binary fixtures."""
    from PIL import Image

    def make(color=(200, 30, 30), size=(64, 48)):
        buf = io.BytesIO()
        Image.new("RGB", size, color).save(buf, format="PNG")
        return buf.getvalue()

    return make


@pytest.fixture
def make_xlsx(tmp_path, png_bytes):
    """Writes a workbook with images anchored at given cells.

    `sheets` maps a sheet title to a list of (cell, colour) pairs.
    """
    import openpyxl
    from openpyxl.drawing.image import Image as XLImage

    def build(sheets=None, name="book.xlsx"):
        sheets = sheets or {"Sheet1": [("C7", (200, 30, 30))]}
        workbook = openpyxl.Workbook()
        default = workbook.active
        first = True
        for title, placements in sheets.items():
            sheet = default if first else workbook.create_sheet()
            sheet.title = title
            first = False
            for cell, color in placements:
                sheet.add_image(XLImage(io.BytesIO(png_bytes(color=color))), cell)
        path = tmp_path / name
        workbook.save(str(path))
        return path

    return build


@pytest.fixture
def make_docx(tmp_path, png_bytes):
    """Writes a Word document with a picture in a table cell and one in a body paragraph."""
    import docx
    from docx.shared import Inches

    def build(table_color=(30, 120, 200), body_color=(20, 200, 90), name="doc.docx",
              with_table=True, with_body=True):
        document = docx.Document()
        document.add_paragraph("Intro paragraph, no picture.")

        if with_table:
            table = document.add_table(rows=2, cols=3)
            run = table.cell(1, 2).paragraphs[0].add_run()
            run.add_picture(io.BytesIO(png_bytes(color=table_color)), width=Inches(1))

        if with_body:
            document.add_paragraph().add_run().add_picture(
                io.BytesIO(png_bytes(color=body_color)), width=Inches(1)
            )

        path = tmp_path / name
        document.save(str(path))
        return path

    return build


@pytest.fixture
def make_pdf(tmp_path, png_bytes):
    """Writes a PDF with one image per page at a known rectangle."""
    import fitz

    def build(pages=None, name="doc.pdf"):
        pages = pages or [(60, 80, 160, 180, (200, 30, 30))]
        doc = fitz.open()
        for x0, y0, x1, y1, color in pages:
            page = doc.new_page()
            page.insert_image(fitz.Rect(x0, y0, x1, y1), stream=png_bytes(color=color))
        path = tmp_path / name
        doc.save(str(path))
        doc.close()
        return path

    return build
