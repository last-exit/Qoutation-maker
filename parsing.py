"""File parsing helpers: rate/unit/venue/date extraction from quotation Excel and PDF files."""
import json
import re
from datetime import datetime
from pathlib import Path

import datefinder
import openpyxl

import image_tools

# --- Rate parsing -----------------------------------------------------------

def clean_rate(val):
    """Parses raw pricing strings into numerical floats, filtering out dimension blocks and units."""
    if val is None:
        return 0.0
    val_str = str(val).strip()
    val_str_lower = val_str.lower()

    for cur in ['aed', 'dhs', 'dh', 'usd', 'aed.', 'dhs.']:
        val_str_lower = val_str_lower.replace(cur, '')

    val_str_lower = val_str_lower.strip()

    # If any alphabetic characters remain, this is not a rate (e.g. "L3 x W3.6m", "Nos")
    if re.search(r'[a-zA-Z]', val_str_lower):
        return 0.0

    val_str = re.sub(r'[^\d.,]', '', val_str)
    if not val_str:
        return 0.0

    if ',' in val_str and '.' in val_str:
        val_str = val_str.replace(',', '')
    elif ',' in val_str:
        parts = val_str.split(',')
        if len(parts[-1]) == 3:
            val_str = val_str.replace(',', '')
        else:
            val_str = val_str.replace(',', '.')

    try:
        return float(val_str)
    except ValueError:
        return 0.0


# --- Unit normalization -------------------------------------------------------

CANONICAL_UNITS = ["Pcs", "Set", "Lump Sum", "Sqm", "Days", "Hrs", "Nos"]

_UNIT_RULES = [
    (r'sq\s*\.?\s*m|sqm|m2|m\^2|square\s*met', "Sqm"),
    (r'\bday', "Days"),
    (r'\bhr|\bhour', "Hrs"),
    (r'\blump|\bl\.?s\.?\b|lumpsum', "Lump Sum"),
    (r'\bset', "Set"),
    (r'\bnos?\b|\bnumber', "Nos"),
    (r'\bpc|\bpiece|\beach|\bea\b|\bqty\b|\bunit\b', "Pcs"),
]

_LUMP_SUM_FALLBACK_KEYWORDS = ["service", "package", "rental", "lumpsum", "lump sum"]


def normalize_unit(raw, description=""):
    """Standardizes messy unit strings into a clean operational unit."""
    raw_str = str(raw).strip().lower() if raw is not None else ""

    if not raw_str or raw_str in ("none", "nan", "-", "n/a"):
        desc_lower = (description or "").lower()
        if any(k in desc_lower for k in _LUMP_SUM_FALLBACK_KEYWORDS):
            return "Lump Sum"
        return "Pcs"

    for pattern, canonical in _UNIT_RULES:
        if re.search(pattern, raw_str):
            return canonical

    # Nonsensical (pure symbols/numbers) unit text
    if not re.search(r'[a-zA-Z]', raw_str):
        desc_lower = (description or "").lower()
        if any(k in desc_lower for k in _LUMP_SUM_FALLBACK_KEYWORDS):
            return "Lump Sum"
        return "Pcs"

    return "Pcs"


# --- Date extraction ----------------------------------------------------------

def extract_date(filename, file_mtime):
    """Extracts date from filename using Regex and datefinder. Fallback to mtime."""
    match_dmy = re.search(r'(\d{1,2})[-_](\d{1,2})[-_](\d{4})', filename)
    if match_dmy:
        return f"{match_dmy.group(3)}-{match_dmy.group(2).zfill(2)}-{match_dmy.group(1).zfill(2)}"

    match_ymd = re.search(r'(\d{4})[-_](\d{1,2})[-_](\d{1,2})', filename)
    if match_ymd:
        return f"{match_ymd.group(1)}-{match_ymd.group(2).zfill(2)}-{match_ymd.group(3).zfill(2)}"

    try:
        matches = list(datefinder.find_dates(filename))
        if matches:
            return matches[0].strftime("%Y-%m-%d")
    except Exception:
        pass

    try:
        dt = datetime.fromtimestamp(file_mtime)
        return dt.strftime("%Y-%m-%d")
    except Exception:
        return "2025-01-01"


# --- Venue extraction -----------------------------------------------------------

VENUES_CONFIG_PATH = Path(__file__).resolve().parent / "venues.json"

_DEFAULT_VENUES = [
    "Kite Beach", "Knowledge Village", "Dubai Mall", "Expo City", "Expo 2020",
    "JBR", "Jumeirah Beach Residence", "La Mer", "Dubai Hills", "City Walk",
    "Global Village", "Downtown Dubai", "Business Bay", "Dubai Marina",
    "Palm Jumeirah", "Bluewaters", "Al Barari", "Zabeel Park",
    "Dubai Festival City", "Al Quoz", "DIFC", "Jumeirah", "Media City",
    "Internet City", "Academic City", "Sports City", "Motor City",
    "Studio City", "Silicon Oasis", "Al Seef", "Madinat Jumeirah",
    "Burj Al Arab", "Atlantis", "World Trade Centre", "WTC", "Yas Island",
    "Abu Dhabi", "Sharjah", "Ajman",
]


def _load_known_venues():
    """Loads the venue keyword list from venues.json so it can be edited without a code change.
    Seeds the file with sane defaults on first run if it doesn't exist yet."""
    try:
        if VENUES_CONFIG_PATH.exists():
            with open(VENUES_CONFIG_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
            venues = data.get("venues", [])
            if venues:
                return venues
    except Exception as e:
        print(f"Failed to load venues.json, using built-in defaults: {e}")

    try:
        with open(VENUES_CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump({"venues": _DEFAULT_VENUES}, f, indent=2)
    except Exception as e:
        print(f"Could not write venues.json: {e}")

    return _DEFAULT_VENUES


KNOWN_VENUES = _load_known_venues()

_VENUE_LABEL_RE = re.compile(r'(?:venue|location|event\s*venue|site)\s*[:\-]\s*(.+)', re.IGNORECASE)
_VENUE_LABEL_CELL_RE = re.compile(r'^(venue|location|event\s*venue|site)$', re.IGNORECASE)


def _extract_labeled_venue_from_text(header_texts):
    """Tier 1: explicit 'Venue: X' / 'Location: X' style labels inside a single text blob
    (works for PDF header lines and single-cell Excel headers)."""
    for h in header_texts:
        if not h:
            continue
        m = _VENUE_LABEL_RE.search(str(h))
        if m:
            candidate = m.group(1).strip()
            candidate = re.split(r'\s{2,}|\t', candidate)[0].strip(" -:|")
            if candidate and 2 <= len(candidate) <= 60 and not re.match(r'^\d+$', candidate):
                return candidate
    return None


def _extract_labeled_venue_from_rows(header_rows):
    """Tier 1b: adjacent-cell 'Venue' | 'Kite Beach' pattern common in Excel header rows,
    where the label and value sit in separate cells rather than one string."""
    for row in header_rows:
        for i, cell in enumerate(row):
            cell_str = str(cell).strip() if cell is not None else ""
            if _VENUE_LABEL_CELL_RE.match(cell_str):
                for nxt in row[i + 1:]:
                    nxt_str = str(nxt).strip() if nxt is not None else ""
                    if nxt_str and not re.match(r'^\d+(\.\d+)?$', nxt_str):
                        return nxt_str
    return None


def extract_venue(filename, header_texts=None, header_rows=None):
    """Extracts venue name from filename and/or scanned header rows.

    Returns (venue, confidence, reason) where confidence is one of
    "high" (explicit labeled field), "medium" (known-venue keyword match),
    "low" (unverified filename/subtitle heuristic), or "none" (nothing found).
    """
    header_texts = header_texts or []
    header_rows = header_rows or []

    labeled = _extract_labeled_venue_from_rows(header_rows) or _extract_labeled_venue_from_text(header_texts)
    if labeled:
        return labeled.title() if labeled.isupper() else labeled, "high", None

    combined = " ".join([filename] + [str(h) for h in header_texts if h])
    combined_lower = combined.lower()

    for venue in KNOWN_VENUES:
        if venue.lower() in combined_lower:
            return venue, "medium", None

    # Text following "@" in the filename (e.g. "... @ Knowledge Village - 9-12-2025.xlsx").
    # Treated as "medium", not a guess: across the real archive the "@" is used deliberately
    # and consistently by the business to mark the site, so demoting it to low-confidence
    # queued ~70 correctly-extracted items for review that needed no attention.
    at_match = re.search(r'@\s*([A-Za-z][A-Za-z0-9 &\'-]{2,40})', filename)
    if at_match:
        candidate = at_match.group(1).strip()
        candidate = re.split(r'\s*-\s*\d', candidate)[0].strip()
        candidate = re.sub(r'\.(pdf|xlsx|xls|docx)$', '', candidate, flags=re.IGNORECASE).strip()
        if candidate:
            return candidate.title(), "medium", None

    # Fallback: trailing segment of a "TITLE - VENUE" style subtitle line
    for h in header_texts:
        h_str = str(h).strip()
        if not h_str or len(h_str) > 80:
            continue
        parts = [p.strip() for p in h_str.split(" - ") if p.strip()]
        if len(parts) >= 2:
            last = parts[-1]
            if not re.search(r'\d', last) and 2 <= len(last) <= 40:
                return last.title(), "low", "venue guessed from subtitle line - please verify"

    return "Venue Unspecified", "none", "no venue signal found in file"


# --- Item assembly ---------------------------------------------------------------

# A product name has to actually name a product. The PDF reader walks raw text lines, and
# anything left standing became a line item — so prices ("2,129.60"), totals ("-2,374.77"),
# dates ("2026-07-24") and block labels ("Client Name:", "Valid Until:") were all indexed as
# products and then surfaced in the Smart Matcher as real, priced results.
#
# The old guard was `^\d+(\.\d+)?$`, which only catches a bare integer or decimal — every
# figure with a thousands separator, currency symbol, sign or bracket walked straight past it.
#
# Deliberately conservative: it rejects only lines that cannot be a product name. Verified
# against the live index — removes 26 bogus rows while keeping every genuine short name
# ("TV Wall", "Lego Table", "Zipline", "Mud Kitchen", "Jungle").
_NUMERIC_ONLY_RE = re.compile(
    r'^[\s$€£]*[-+(]?[\d,.\s]+%?\)?\s*(aed|usd|eur|dhs?|dirhams?)?\s*[-+)]?\s*$', re.I)
_DATE_ONLY_RE = re.compile(r'^\s*\d{1,4}\s*[-/.]\s*\d{1,2}\s*[-/.]\s*\d{1,4}\s*$')

# Layout words that survive as their own "line" when a table cell wraps in the PDF.
_FRAGMENT_WORDS = {
    'only', 'as req', 'as required', 'items', 'item', 'each', 'req', 'note', 'notes',
    'incl', 'included', 'excl', 'etc', 'and', 'the', 'for', 'with', 'na', 'n/a',
    'tbc', 'tbd', 'yes', 'no', 'all', 'new', 'same', 'total', 'subtotal',
}


def looks_like_description(text):
    """True when a text line could plausibly be a product/service name."""
    t = str(text or "").strip()
    if len(t) < 3:
        return False
    if _NUMERIC_ONLY_RE.match(t) or _DATE_ONLY_RE.match(t):
        return False
    if sum(ch.isalpha() for ch in t) < 3:
        return False
    core = t.rstrip(':').strip().lower()
    if core in _FRAGMENT_WORDS:
        return False
    # "Client Name:", "Valid Until:", "Timber House:" — a short trailing-colon line is the
    # label of a block, not the thing being sold.
    if t.endswith(':') and len(core.split()) <= 2:
        return False
    return True


# Lines that read as continuation detail for the item above them rather than a new item:
# dimension callouts, bulleted features, and parenthetical notes.
_SPEC_LINE_RE = re.compile(
    r'^\s*([-•*•]|\(|[LWHD]\s*\d|\d+\s*(mm|cm|m|sqm|ft)\b|measures\b|features\b|components\b|includes?\b)',
    re.I)


def looks_like_spec_line(text):
    """True when a line is specification detail belonging to the preceding item."""
    t = str(text or "").strip()
    if not t or len(t) > 160:
        return False
    if _NUMERIC_ONLY_RE.match(t) or _DATE_ONLY_RE.match(t):
        return False
    return bool(_SPEC_LINE_RE.match(t)) or bool(re.search(r'\d+\s*[xX]\s*\d+', t))


# Markers written by the PM actions in the review queue. Once present, an item is never
# re-flagged, so approving/dismissing survives a re-sync of the same source file.
REVIEW_CLEARED_MARKERS = ("corrected by pm", "dismissed by pm")


def evaluate_review_flags(meta):
    """Decides whether an indexed item needs PM review, returning (needs_review, reasons).

    Scoped deliberately to *actionable data gaps* a PM can actually fix — a missing rate,
    an unverified venue, a missing description, or a rate that failed to reconcile against
    its own line total. Parser-confidence warnings that fire for a whole file at a time
    (e.g. "rate column guessed heuristically") are recorded on the item but do NOT by
    themselves queue it: doing so flagged 100% of the index and made the queue useless
    as a worklist.

    Evaluated at read time rather than trusted from stored metadata, so tightening these
    rules takes effect immediately without forcing a full re-index.
    """
    prior_reason = str(meta.get("flag_reason", "")).lower()
    if any(marker in prior_reason for marker in REVIEW_CLEARED_MARKERS):
        return False, []

    reasons = []

    try:
        rate = float(meta.get("historical_rate", 0) or 0)
    except (TypeError, ValueError):
        rate = 0.0
    if rate <= 0:
        reasons.append("missing unit rate (0.00 AED)")

    description = str(meta.get("original_description", "") or "").strip()
    if len(description) < 3:
        reasons.append("missing description")

    venue = str(meta.get("venue", "") or "").strip()
    if not venue or venue == "Venue Unspecified" or meta.get("venue_confidence") in ("low", "none"):
        reasons.append("unverified venue")

    if "reconcile" in prior_reason:
        reasons.append("rate did not reconcile against the line total")

    return bool(reasons), reasons


def _build_item(description, rate, unit, quote_date, venue, venue_confidence, venue_reason,
                 file_name, image_base64, rate_confidence, rate_reason):
    """Assembles a parsed line item with confidence metadata attached, so low-confidence
    extractions can be routed to a review queue instead of silently entering the index."""
    reasons = [r for r in (rate_reason, venue_reason) if r]
    needs_review, _ = evaluate_review_flags({
        "historical_rate": rate,
        "original_description": description,
        "venue": venue,
        "venue_confidence": venue_confidence,
        "flag_reason": "; ".join(reasons),
    })
    return {
        'original_description': description,
        'historical_rate': rate,
        'unit': unit,
        'quote_date': quote_date,
        'venue': venue,
        'file_name': file_name,
        'image_base64': image_base64,
        'rate_confidence': rate_confidence,
        'venue_confidence': venue_confidence,
        'needs_review': needs_review,
        'flag_reason': '; '.join(reasons),
    }


# --- Column classification ----------------------------------------------------

_RATE_EXCLUDE_KEYWORDS = ["total", "vat", "amount", "grand", "sub total", "subtotal", "tax"]
_RATE_KEYWORDS_HIGH = ["unit price", "unit cost", "unit rate", "rate"]
_RATE_KEYWORDS_LOW = ["price", "cost"]


def classify_columns(row_str):
    """Given a header row (list of lowercased cell strings), classify column indices strictly.

    Never lets a Total/VAT/Amount column become the rate column, and prefers exact
    "rate"/"unit price" headers over generic "price" matches.
    """
    col_map = {}
    rate_candidates = []

    for idx, val in enumerate(row_str, start=1):
        if not val:
            continue

        if any(k in val for k in _RATE_EXCLUDE_KEYWORDS):
            col_map.setdefault("total", idx)
            continue

        if "description" in val or "particulars" in val or "scope of work" in val:
            col_map.setdefault("description", idx)
        elif "image" in val or "photo" in val or "picture" in val:
            col_map.setdefault("images", idx)
        elif "item #" in val or "item no" in val or "s.no" in val or "sr no" in val or "sl no" in val:
            col_map.setdefault("item_num", idx)
        elif "qty" in val or "quantity" in val:
            col_map.setdefault("qty", idx)
        elif any(k in val for k in _RATE_KEYWORDS_HIGH):
            rate_candidates.append((0, idx))
        elif any(k in val for k in _RATE_KEYWORDS_LOW):
            rate_candidates.append((1, idx))
        elif val.strip() == "unit" or val.strip().startswith("unit"):
            col_map.setdefault("unit", idx)

    if rate_candidates:
        rate_candidates.sort(key=lambda t: t[0])
        col_map["rate"] = rate_candidates[0][1]

    return col_map


def _heuristic_columns(sheet, sample_rows=15):
    """Conservative fallback column detection when no keyword-based header row is found.

    Only returns a result when it can positively validate description (mostly text)
    and rate (mostly plausible positive numbers, distinct from a likely qty column).
    Returns {} (skip the sheet) rather than guessing wrong.
    """
    max_col = min(sheet.max_column or 1, 15)
    col_stats = {}

    for r in range(1, sample_rows + 1):
        for c in range(1, max_col + 1):
            val = sheet.cell(row=r, column=c).value
            if val is None:
                continue
            stats = col_stats.setdefault(c, {"text_len": 0, "text_count": 0, "num_count": 0, "small_int": 0})
            if isinstance(val, (int, float)):
                stats["num_count"] += 1
                if 0 < val <= 500 and float(val) == int(val):
                    stats["small_int"] += 1
            else:
                s = str(val).strip()
                if s and not re.match(r'^[\d.,\s]+$', s):
                    stats["text_len"] += len(s)
                    stats["text_count"] += 1

    if not col_stats:
        return {}

    desc_col = max(col_stats, key=lambda c: col_stats[c]["text_len"])
    if col_stats[desc_col]["text_count"] < 3:
        return {}

    numeric_cols = sorted([c for c in col_stats if c != desc_col and col_stats[c]["num_count"] >= 2])
    if not numeric_cols:
        return {}

    qty_col = min(numeric_cols, key=lambda c: (-col_stats[c]["small_int"], c))
    rate_candidates = [c for c in numeric_cols if c != qty_col]
    if not rate_candidates:
        return {}

    # Prefer the rate candidate furthest right that isn't the very last (likely "Total") column,
    # since layouts are typically Desc..Unit..Qty..Rate..VAT..Total.
    rate_candidates.sort()
    rate_col = rate_candidates[-2] if len(rate_candidates) >= 2 else rate_candidates[0]

    return {"description": desc_col, "qty": qty_col, "rate": rate_col}


# --- Excel parsing --------------------------------------------------------------

def parse_excel_file(file_path):
    """Parses Excel sheets dynamically: strict rate-column detection, venue/date/unit metadata."""
    items = []
    path_obj = Path(file_path)
    file_name = path_obj.name

    try:
        mtime = path_obj.stat().st_mtime
        file_date = extract_date(file_name, mtime)
    except Exception:
        file_date = "2025-01-01"

    try:
        wb = openpyxl.load_workbook(str(path_obj), data_only=True)
    except Exception as e:
        print(f"Failed to open Excel '{file_name}': {e}")
        return items

    active_sheet_title = wb.active.title
    for sheet in wb.worksheets:
        sheet_title_lower = sheet.title.lower()
        should_parse = (
            sheet.title == active_sheet_title or
            any(x in sheet_title_lower for x in ["option", "opt", "sheet", "quotation", "production"])
        )
        if not should_parse:
            continue

        header_row_idx = None
        col_map = {}
        header_texts = []
        header_rows = []
        used_heuristic_columns = False

        for r_idx in range(1, 16):
            try:
                row_vals = [sheet.cell(row=r_idx, column=c_idx).value for c_idx in range(1, 16)]
            except Exception:
                continue
            for v in row_vals:
                if v is not None and str(v).strip():
                    header_texts.append(str(v).strip())
            header_rows.append(row_vals)
            row_str = [str(val).strip().lower() if val is not None else "" for val in row_vals]

            if any("description" in v or "particulars" in v for v in row_str):
                header_row_idx = r_idx
                col_map = classify_columns(row_str)
                break

        if "description" not in col_map or "rate" not in col_map:
            heuristic = _heuristic_columns(sheet)
            if not heuristic:
                # Could not confidently locate columns — skip this sheet rather than guess wrong.
                continue
            col_map = {**heuristic, **col_map}
            used_heuristic_columns = True
            if header_row_idx is None:
                header_row_idx = 1

        desc_col = col_map.get("description")
        rate_col = col_map.get("rate")
        unit_col = col_map.get("unit")
        img_col = col_map.get("images")
        qty_col = col_map.get("qty")
        total_col = col_map.get("total")

        if not desc_col or not rate_col:
            continue

        # Map embedded product images to their rows (0-indexed anchor rows). Prefer an exact
        # match against a detected "Images" column header; but real cost sheets often embed
        # photos without ever labeling a column for them, so if no image column was found,
        # fall back to matching by row position alone rather than silently dropping every photo.
        images_by_row = {}
        if hasattr(sheet, '_images') and sheet._images:
            for img in sheet._images:
                try:
                    col_idx = img.anchor._from.col
                    row_idx = img.anchor._from.row
                    if img_col:
                        if col_idx == (img_col - 1):
                            images_by_row[row_idx + 1] = img
                    else:
                        images_by_row[row_idx + 1] = img
                except Exception as e:
                    print(f"Error mapping image: {e}")

        start_row = (header_row_idx + 1) if header_row_idx else 2
        for r in range(start_row, sheet.max_row + 1):
            try:
                desc_val = sheet.cell(row=r, column=desc_col).value
                rate_val = sheet.cell(row=r, column=rate_col).value
                unit_val = sheet.cell(row=r, column=unit_col).value if unit_col else None
                qty_val = sheet.cell(row=r, column=qty_col).value if qty_col else None
                total_val = sheet.cell(row=r, column=total_col).value if total_col else None
            except Exception:
                continue

            if desc_val is None or str(desc_val).strip() == "":
                continue

            desc_str = str(desc_val).strip()
            desc_lower = desc_str.lower()
            if "total" in desc_lower or "subtotal" in desc_lower or "grand total" in desc_lower:
                continue

            rate_float = clean_rate(rate_val)
            qty_float = clean_rate(qty_val) if qty_val is not None else 0.0
            total_float = clean_rate(total_val) if total_val is not None else 0.0

            if rate_float <= 0:
                # A priced-out row whose rate cell is blank is a real data gap the PM should
                # see in the review queue, not something to silently drop. Require corroborating
                # evidence that this is actually a line item (a qty, unit or total) so section
                # headings and spacer rows don't flood the queue.
                looks_like_line_item = (
                    qty_float > 0 or total_float > 0 or
                    (unit_val is not None and str(unit_val).strip() != "")
                )
                if not looks_like_line_item:
                    continue

            # Strict validation: if both Qty and a Total column are present and populated,
            # cross-check the rate is consistent (roughly total / qty), catching cases where
            # the wrong column was picked up as "rate".
            rate_validated = False
            rate_reconcile_failed = False
            if rate_float > 0 and qty_float > 0 and total_float > 0:
                expected_rate = total_float / qty_float
                # Allow generous tolerance for VAT-inclusive totals (up to ~10%) and rounding.
                if 0.5 * expected_rate <= rate_float <= 1.3 * expected_rate:
                    rate_validated = True
                else:
                    rate_reconcile_failed = True

            if rate_float <= 0:
                rate_confidence, rate_reason = "low", "rate cell is blank or zero in the source file"
            elif rate_reconcile_failed:
                # Kept rather than dropped: a mismatch means the number is suspect, and the PM
                # can only fix what they can see. It surfaces in the review queue.
                rate_confidence, rate_reason = "low", "rate did not reconcile against the line total - verify"
            elif rate_validated:
                rate_confidence, rate_reason = "high", None
            elif used_heuristic_columns:
                rate_confidence, rate_reason = "low", "rate column guessed heuristically (no header keywords found in this sheet)"
            else:
                rate_confidence, rate_reason = "medium", "rate column matched by header keyword but not cross-checked (no qty/total to verify against)"

            unit_str = normalize_unit(unit_val, desc_str)

            img_base64 = ""
            if r in images_by_row:
                img_base64 = image_tools.get_embedded_image_base64(images_by_row[r])

            venue, venue_confidence, venue_reason = extract_venue(file_name, header_texts, header_rows)

            items.append(_build_item(
                description=desc_str, rate=rate_float, unit=unit_str, quote_date=file_date,
                venue=venue, venue_confidence=venue_confidence, venue_reason=venue_reason,
                file_name=file_name, image_base64=img_base64,
                rate_confidence=rate_confidence, rate_reason=rate_reason,
            ))

    return items


# --- Word (.docx) parsing ----------------------------------------------------------

def _distinct_row_cells(row):
    """Returns a row's genuinely distinct cells.

    python-docx repeats the same cell object for horizontally merged spans, so a 5-column
    row can report 8 cells with duplicated text. De-duplicating on the underlying <tc>
    element keeps column indices aligned with the header row.
    """
    seen, out = set(), []
    for cell in row.cells:
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        out.append(cell)
    return out


def _cell_image_base64(cell, document):
    """Pulls the first embedded image out of a table cell, thumbnailed like every other
    image source, so photos sitting inside Word quotation tables reach the index too."""
    try:
        from docx.oxml.ns import qn
        for blip in cell._tc.findall('.//' + qn('a:blip')):
            rid = blip.get(qn('r:embed')) or blip.get(qn('r:link'))
            if not rid:
                continue
            part = document.part.related_parts.get(rid)
            if part is None:
                continue
            return image_tools.bytes_to_thumbnail_base64(part.blob)
    except Exception as e:
        print(f"Failed to extract image from docx cell: {e}")
    return ""


def parse_docx_file(file_path):
    """Extracts priced line items (and their embedded photos) from Word quotation tables.

    Several documents in the live archive exist only as .docx, and were previously not
    indexed at all, so both their pricing and their product photos were invisible to the app.
    """
    items = []
    path_obj = Path(file_path)
    file_name = path_obj.name

    try:
        file_date = extract_date(file_name, path_obj.stat().st_mtime)
    except Exception:
        file_date = "2025-01-01"

    try:
        from docx import Document
        document = Document(str(path_obj))
    except Exception as e:
        print(f"Failed to open Word document '{file_name}': {e}")
        return items

    header_texts = [p.text.strip() for p in document.paragraphs[:20] if p.text.strip()]

    for table in document.tables:
        header_idx, col_map = None, {}

        for r_idx, row in enumerate(table.rows[:6]):
            cells = _distinct_row_cells(row)
            row_str = [c.text.strip().lower() for c in cells]
            header_texts.extend(c.text.strip() for c in cells if c.text.strip())
            if any(("description" in v or "item" in v or "particulars" in v) for v in row_str) and \
               any(("cost" in v or "rate" in v or "price" in v or "amount" in v) for v in row_str):
                header_idx = r_idx
                col_map = classify_columns(row_str)
                # "Item" is this archive's usual description header but is too generic for the
                # shared classifier, so it's resolved here where the table context is known.
                if "description" not in col_map:
                    for i, v in enumerate(row_str, start=1):
                        if v.strip() in ("item", "items", "item name", "element"):
                            col_map["description"] = i
                            break
                break

        if header_idx is None or "description" not in col_map:
            continue

        desc_col = col_map["description"]
        rate_col = col_map.get("rate")
        qty_col = col_map.get("qty")
        unit_col = col_map.get("unit")
        total_col = col_map.get("total")

        venue, venue_confidence, venue_reason = extract_venue(file_name, header_texts)

        for row in table.rows[header_idx + 1:]:
            cells = _distinct_row_cells(row)

            def cell_text(col):
                return cells[col - 1].text.strip() if col and 0 < col <= len(cells) else ""

            desc_str = cell_text(desc_col)
            if not desc_str or len(desc_str) < 3:
                continue
            if any(k in desc_str.lower() for k in ("total", "subtotal", "grand total")):
                continue

            rate_float = clean_rate(cell_text(rate_col)) if rate_col else 0.0
            qty_float = clean_rate(cell_text(qty_col)) if qty_col else 0.0
            total_float = clean_rate(cell_text(total_col)) if total_col else 0.0

            # Merged/blank rate cells are common here; fall back to deriving it from the total.
            if rate_float <= 0 and total_float > 0 and qty_float > 0:
                rate_float = total_float / qty_float

            if rate_float <= 0 and total_float <= 0:
                continue

            if rate_float <= 0:
                rate_confidence, rate_reason = "low", "rate cell is blank or zero in the source file"
            elif qty_float > 0 and total_float > 0:
                expected = total_float / qty_float
                if 0.5 * expected <= rate_float <= 1.3 * expected:
                    rate_confidence, rate_reason = "high", None
                else:
                    rate_confidence, rate_reason = "low", "rate did not reconcile against the line total - verify"
            else:
                rate_confidence, rate_reason = "medium", "rate read from Word table without a total to verify against"

            image_b64 = ""
            for cell in cells:
                image_b64 = _cell_image_base64(cell, document)
                if image_b64:
                    break

            items.append(_build_item(
                description=desc_str, rate=rate_float,
                unit=normalize_unit(cell_text(unit_col) if unit_col else None, desc_str),
                quote_date=file_date, venue=venue, venue_confidence=venue_confidence,
                venue_reason=venue_reason, file_name=file_name, image_base64=image_b64,
                rate_confidence=rate_confidence, rate_reason=rate_reason,
            ))

    return items


# --- PDF parsing ------------------------------------------------------------------

def _get_pdf_lines_with_bbox(page):
    """Extracts text lines the same way page.get_text('text') does, but keeps each line's
    vertical position too, so a line item can be matched to a nearby embedded photo.
    Verified to produce identical line splitting to the plain-text path on representative
    layouts — this is a superset of the old extraction, not a behavior change."""
    lines, boxes = [], []
    try:
        raw = page.get_text("dict")
        for block in raw.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                text = "".join(span.get("text", "") for span in line.get("spans", [])).strip()
                if not text:
                    continue
                bbox = line.get("bbox")
                if not bbox:
                    spans_bboxes = [s.get("bbox") for s in line.get("spans", []) if s.get("bbox")]
                    if spans_bboxes:
                        bbox = (min(b[0] for b in spans_bboxes), min(b[1] for b in spans_bboxes),
                                max(b[2] for b in spans_bboxes), max(b[3] for b in spans_bboxes))
                    else:
                        bbox = (0, 0, 0, 0)
                lines.append(text)
                boxes.append((bbox[1], bbox[3]))  # (y0, y1)
    except Exception as e:
        print(f"Failed to extract positioned text from PDF page: {e}")
    return lines, boxes


def _extract_pdf_page_images(page, max_images=40):
    """Pulls every embedded image on a page with its vertical position, thumbnailed and
    normalized the same way Excel-embedded images are (see image_tools.get_embedded_image_base64),
    so photos already sitting in old quotation PDFs actually make it into the historical index
    instead of every PDF item silently having no image."""
    results = []
    try:
        doc = page.parent
        for info in page.get_image_info(xrefs=True)[:max_images]:
            xref = info.get("xref")
            bbox = info.get("bbox")
            if not xref or not bbox:
                continue
            try:
                img_dict = doc.extract_image(xref)
                raw_bytes = img_dict.get("image")
                if not raw_bytes:
                    continue
                thumb_b64 = image_tools.bytes_to_thumbnail_base64(raw_bytes)
                results.append({"y0": bbox[1], "y1": bbox[3], "image_base64": thumb_b64, "used": False})
            except Exception as e:
                print(f"Failed to extract embedded PDF image (xref {xref}): {e}")
    except Exception as e:
        print(f"Failed to read image info from PDF page: {e}")
    return results


def assign_images_to_rows(row_positions, page_images):
    """Assigns each page image to the table row it visually belongs to.

    In these quotation PDFs a product photo is a tall graphic spanning its whole table row
    (e.g. y 190-311) while the row's first text line sits near the top of that band
    (y 170-182). Comparing centre points therefore put them ~75pt apart and a fixed
    tolerance rejected the match, which is why only a fraction of the embedded photos ever
    reached the index.

    Instead each row is given a band running from its own first line down to the start of
    the next row, and every image is awarded to the row band it overlaps most. Rows are
    defined by the items actually extracted, so this stays correct regardless of row height,
    photo size, or how many lines of description a row wraps to.

    row_positions: list of (row_index, y_top) for each item found on the page, in order.
    Returns {row_index: image_base64}.
    """
    if not page_positions_valid(row_positions) or not page_images:
        return {}

    ordered = sorted(row_positions, key=lambda rp: rp[1])
    bands = []
    for i, (row_idx, y_top) in enumerate(ordered):
        y_bottom = ordered[i + 1][1] if i + 1 < len(ordered) else float('inf')
        # Start slightly above the text line: a row's photo often begins fractionally
        # higher than its first character.
        bands.append((row_idx, y_top - 6.0, y_bottom))

    assigned = {}
    for img in page_images:
        if img["used"]:
            continue
        best_row, best_overlap = None, 0.0
        for row_idx, band_top, band_bottom in bands:
            if row_idx in assigned:
                continue
            overlap = min(img["y1"], band_bottom) - max(img["y0"], band_top)
            if overlap > best_overlap:
                best_row, best_overlap = row_idx, overlap
        if best_row is not None and best_overlap > 0:
            assigned[best_row] = img["image_base64"]
            img["used"] = True

    return assigned


def page_positions_valid(row_positions):
    return bool(row_positions)


def _extract_trailing_numbers(line, max_count=4):
    """Pulls up to max_count clean numeric tokens off the right end of a line
    (handles rows that PyMuPDF flattens into one line, e.g. "Backdrop 10x6m  2  450.00  900.00").
    Returns (numbers_in_left_to_right_order, remaining_description_text)."""
    tokens = [t for t in re.split(r'\s{1,}|\|', line.strip()) if t != '']
    numbers = []
    i = len(tokens) - 1
    while i >= 0 and len(numbers) < max_count:
        val = clean_rate(tokens[i])
        if val > 0:
            numbers.insert(0, val)
            i -= 1
        else:
            break
    remaining = " ".join(tokens[:i + 1]).strip(" -:|")
    return numbers, remaining


def _collect_spec_lines(lines, start, limit=8):
    """Gathers the specification lines that follow an item row in a PDF.

    A PDF gives one text line at a time, so an item's measures and features arrive as separate
    lines after its name. They used to be either dropped or — worse — picked up by the next
    loop iteration and indexed as items in their own right. Collecting them here means the
    spec travels with the product, which is what the client actually reads on the quote.

    Stops at the first line that looks like a new item so it can never swallow the next row.
    """
    spec = []
    for offset in range(limit):
        i = start + offset
        if i >= len(lines):
            break
        candidate = str(lines[i]).strip()
        if not candidate:
            continue
        if not looks_like_spec_line(candidate):
            break
        spec.append(candidate)
    return spec


def parse_pdf_file(file_path):
    """Extracts items from PDF text via PyMuPDF (fitz), with venue/unit normalization.

    Tries same-line trailing-numbers detection first (handles cleanly-flattened table rows,
    and can cross-validate qty*rate against a total like the Excel parser does). Falls back to
    the older multi-line positional lookahead for messier layouts, but tags that path as
    low-confidence since it's guessing at column order.
    """
    items = []
    path_obj = Path(file_path)
    file_name = path_obj.name

    try:
        mtime = path_obj.stat().st_mtime
        file_date = extract_date(file_name, mtime)
    except Exception:
        file_date = "2025-01-01"

    try:
        import fitz
        doc = fitz.open(str(path_obj))
    except Exception as e:
        print(f"Failed to open PDF '{file_name}': {e}")
        return items

    header_texts = []
    venue, venue_confidence, venue_reason = "Venue Unspecified", "none", "no venue signal found in file"

    for page in doc:
        try:
            lines, boxes = _get_pdf_lines_with_bbox(page)
        except Exception as e:
            print(f"Failed to extract text from PDF: {e}")
            continue

        page_images = _extract_pdf_page_images(page)
        # Row-band image assignment needs every row on the page before it can decide, so
        # positions are collected during the scan and images attached in a pass afterwards.
        page_row_positions = []

        if not header_texts:
            header_texts = lines[:20]
            venue, venue_confidence, venue_reason = extract_venue(file_name, header_texts)

        idx = 0
        while idx < len(lines):
            line = lines[idx]

            line_lower = line.lower()
            if len(line) < 4 or any(x in line_lower for x in [
                "page ", "quotation #", "date:", "client:", "venue:", "location:", "site:", "attn", "attention",
            ]):
                idx += 1
                continue

            if "total" in line_lower or "subtotal" in line_lower:
                idx += 1
                continue

            if re.match(r'^\d+(\.\d+)?$', line):
                idx += 1
                continue

            # Table header row (e.g. "Description  Qty  Rate  Amount") — not an item itself.
            if "description" in line_lower and any(k in line_lower for k in ["rate", "price", "amount", "qty", "unit"]):
                idx += 1
                continue

            # Some PDF table layouts place each header cell as its own separate text line
            # (e.g. "Description" and "Rate" never sharing one line) rather than one flattened
            # header row — catch those bare column-label lines too, or a lone "Description"
            # can wrongly get treated as an item and swallow the real row that follows it.
            if line_lower in (
                "description", "item", "item description", "qty", "quantity", "rate",
                "unit price", "unit cost", "unit rate", "price", "cost", "amount",
                "total", "unit", "sr.no", "sr. no", "s.no", "s. no", "no.",
            ):
                idx += 1
                continue

            # All-caps line with no digits reads as a document title / section heading, not a line item —
            # letting these through was the original bug: a title could greedily grab a distant unrelated
            # number in the tier-2 lookahead below and jump straight over the real item rows that followed.
            if line.isupper() and not re.search(r'\d', line):
                idx += 1
                continue

            # --- Tier 1: same-line trailing numbers (row already flattened onto one line) ---
            trailing_numbers, remaining_desc = _extract_trailing_numbers(line)
            if looks_like_description(remaining_desc) and len(trailing_numbers) >= 2:
                if len(trailing_numbers) >= 3:
                    qty, rate, total = trailing_numbers[-3], trailing_numbers[-2], trailing_numbers[-1]
                    expected = qty * rate
                    if qty > 0 and rate > 0 and total > 0 and 0.5 * expected <= total <= 1.3 * expected:
                        rate_confidence, rate_reason = "high", None
                    else:
                        rate_confidence, rate_reason = "low", "same-line qty*rate didn't reconcile with total - verify"
                else:
                    rate = trailing_numbers[-1]
                    rate_confidence, rate_reason = "medium", "same-line column order assumed (rate = last value) - verify"

                if rate > 0:
                    line_y0, _ = boxes[idx] if idx < len(boxes) else (0, 0)
                    page_row_positions.append((len(items), line_y0))
                    spec = _collect_spec_lines(lines, idx + 1)
                    full_desc = remaining_desc + ("\n" + "\n".join(spec) if spec else "")
                    items.append(_build_item(
                        description=full_desc, rate=rate, unit=normalize_unit(None, remaining_desc),
                        quote_date=file_date, venue=venue, venue_confidence=venue_confidence,
                        venue_reason=venue_reason, file_name=file_name, image_base64="",
                        rate_confidence=rate_confidence, rate_reason=rate_reason,
                    ))
                    idx += 1
                    continue

            # --- Tier 2: fallback multi-line positional lookahead (messier / wrapped layouts) ---
            rate_found = None
            unit_found = None
            offset_used = 0

            numbers = []
            for offset in range(1, 15):
                if idx + offset < len(lines):
                    next_line = lines[idx + offset]
                    val = clean_rate(next_line)
                    if val > 0:
                        numbers.append((val, offset))

            if numbers:
                large_numbers = [n for n in numbers if n[0] > 10.0]
                if large_numbers:
                    rate_found = large_numbers[0][0]
                    offset_used = large_numbers[0][1]

                    if offset_used > 1:
                        for u_offset in range(1, offset_used):
                            potential_unit = lines[idx + u_offset]
                            if len(potential_unit) < 8 and re.search(r'[a-zA-Z]', potential_unit):
                                unit_found = potential_unit
                                break
                else:
                    rate_found = numbers[-1][0]
                    offset_used = numbers[-1][1]

            # Tier 2 is a positional guess, so the description gate matters most here: this is
            # the path that turned "2,129.60" and "Client Name:" into priced line items.
            if rate_found is not None and looks_like_description(line):
                line_y0, _ = boxes[idx] if idx < len(boxes) else (0, 0)
                page_row_positions.append((len(items), line_y0))
                spec = _collect_spec_lines(lines, idx + 1)
                full_line = line + ("\n" + "\n".join(spec) if spec else "")
                items.append(_build_item(
                    description=full_line, rate=rate_found, unit=normalize_unit(unit_found, line),
                    quote_date=file_date, venue=venue, venue_confidence=venue_confidence,
                    venue_reason=venue_reason, file_name=file_name, image_base64="",
                    rate_confidence="low",
                    rate_reason="rate located via multi-line positional guess - verify",
                ))
                idx += offset_used + 1
            else:
                idx += 1

        # Attach this page's photos now that every row band on it is known.
        for row_idx, image_b64 in assign_images_to_rows(page_row_positions, page_images).items():
            items[row_idx]['image_base64'] = image_b64

    return items
