"""Dynamic Excel/Word quotation generation.

Rows are generated to exactly match the item list — no leftover blank template rows,
regardless of whether the draft has 1 item or 100.
"""
import base64
import io
import json
from datetime import datetime, timedelta
from pathlib import Path

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.cell.rich_text import CellRichText, TextBlock
from openpyxl.cell.text import InlineFont
from openpyxl.drawing.spreadsheet_drawing import AnchorMarker, OneCellAnchor
from openpyxl.drawing.xdr import XDRPositiveSize2D
from openpyxl.utils.units import pixels_to_EMU
from PIL import Image as PILImage, ImageOps

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import image_store

NAVY_HEX = "1F497D"  # kept as a fallback constant; live styling now reads from COMPANY below

# --- Company branding config -------------------------------------------------------

COMPANY_CONFIG_PATH = Path(__file__).resolve().parent / "company.json"

# Confirmed by the business owner: "Red Cube" is the company running this app; "Boom Tree"
# is one of their clients (their Drive sync folder is just named after that client's
# ongoing job). Edit company.json directly to correct name/colors/logo/PM — no code change
# needed; this dict only seeds the file the first time it's created.
# primary_hex/accent_hex are sampled from the real Red Cube logo (near-black + red) rather
# than an invented palette, so generated documents carry the company's actual brand colors.
_DEFAULT_COMPANY = {
    "name": "RED CUBE",
    "tagline": "Quotation & Cost Estimate",
    "primary_hex": "141313",
    "accent_hex": "DB302F",
    "logo_path": "assets/red_cube_logo.png",
    "pm_name": "",
}


def _load_company_config():
    try:
        if COMPANY_CONFIG_PATH.exists():
            with open(COMPANY_CONFIG_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
            if data.get("name"):
                return {**_DEFAULT_COMPANY, **data}
    except Exception as e:
        print(f"Failed to load company.json, using built-in defaults: {e}")

    try:
        with open(COMPANY_CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(_DEFAULT_COMPANY, f, indent=2)
    except Exception as e:
        print(f"Could not write company.json: {e}")

    return _DEFAULT_COMPANY


COMPANY = _load_company_config()

# --- Terms / validity config ------------------------------------------------------

TERMS_CONFIG_PATH = Path(__file__).resolve().parent / "terms.json"

_DEFAULT_TERMS = {
    "validity_days": 14,
    "payment_terms": [
        "50% advance payment required upon confirmation of order.",
        "Balance 50% payable prior to event build / handover.",
        "Prices are exclusive of any items not explicitly listed above.",
    ],
}


def _load_terms_config():
    """Loads payment terms / default validity period from terms.json so the business can
    edit its own boilerplate without a code change. Seeds the file with defaults on first run."""
    try:
        if TERMS_CONFIG_PATH.exists():
            with open(TERMS_CONFIG_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
            if data.get("payment_terms"):
                return data
    except Exception as e:
        print(f"Failed to load terms.json, using built-in defaults: {e}")

    try:
        with open(TERMS_CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(_DEFAULT_TERMS, f, indent=2)
    except Exception as e:
        print(f"Could not write terms.json: {e}")

    return _DEFAULT_TERMS


TERMS_CONFIG = _load_terms_config()


def compute_valid_until(quote_date_str, validity_days=None):
    """Returns quote_date + validity_days as YYYY-MM-DD. Falls back to today + configured
    default if quote_date_str can't be parsed."""
    days = int(validity_days if validity_days is not None else TERMS_CONFIG.get("validity_days", 14))
    try:
        base_date = datetime.strptime(quote_date_str, "%Y-%m-%d")
    except Exception:
        base_date = datetime.now()
    return (base_date + timedelta(days=days)).strftime("%Y-%m-%d")


def compute_totals(items, discount_type=None, discount_value=0.0):
    """Subtotal -> discount -> VAT on discounted subtotal -> grand total. Shared by
    Excel/Word generation and the history record so numbers always match the UI."""
    subtotal = sum(float(i.get("qty", 0) or 0) * float(i.get("rate", 0) or 0) for i in items)
    discount_value = float(discount_value or 0)

    if discount_type == "percent":
        discount_amount = subtotal * (discount_value / 100.0)
    elif discount_type == "flat":
        discount_amount = discount_value
    else:
        discount_amount = 0.0

    discount_amount = max(0.0, min(discount_amount, subtotal))
    discounted_subtotal = subtotal - discount_amount
    vat = discounted_subtotal * 0.05
    grand_total = discounted_subtotal + vat

    return {
        "subtotal": round(subtotal, 2),
        "discount_amount": round(discount_amount, 2),
        "discounted_subtotal": round(discounted_subtotal, 2),
        "vat": round(vat, 2),
        "grand_total": round(grand_total, 2),
    }


# --- Image handling ---------------------------------------------------------------

THUMB_BOX = 155  # px — display size for line-item photos in the generated Excel sheet

# Pixel resolution the embedded image is rendered at, kept deliberately far above the size it
# is *displayed* at so the photo stays sharp when the document is printed or zoomed. The old
# code rendered at THUMB_BOX (60px) and displayed at 1.4cm, which is ~110 DPI — visibly soft.
# 480px across a 3.4cm print width is ~360 DPI.
DOC_IMAGE_PX = 700
WORD_IMAGE_CM = 4.5  # must stay inside the image column's width, else the table is forced wider
LOGO_TARGET_DPI = 140  # logo print width is derived from its own resolution, see masthead

# --- Shared table palette ----------------------------------------------------------
# Modelled on the business's own cost sheets, where each band of the table is a different
# colour so the eye can separate line items from subtotals from the final figure at a glance.
# Used by BOTH the Excel and Word generators so the two outputs look like one document.
BAND_FILL = "F7F5F1"      # every other line item
SUBTOTAL_FILL = "EFEDE7"  # subtotal / VAT rows
DISCOUNT_FILL = "DCE9F5"  # discount row
DOC_FONT = "Arial"        # professional, present on every Windows/Office install

# A line item's identity is its name *and* its specification. Splitting them lets the name be
# emphasised while the measures/features below stay readable, instead of one flat blob.
def split_name_and_spec(description):
    lines = [ln.rstrip() for ln in str(description or "").split("\n")]
    lines = [ln for ln in lines if ln.strip()]
    if not lines:
        return "", []
    return lines[0].strip(), lines[1:]

# Images that fail to embed are counted here so the caller can tell the PM a photo went
# missing. Previously the only trace was a print() to a console a desktop user never sees:
# the quote reported success and the image was silently absent from the document.
IMAGE_FAILURES = []


def load_image_bytes(value):
    """Resolves a line item's image field to raw bytes.

    Accepts a content-addressed ref (the current form) or a legacy inline
    `data:image/png;base64,...` URI, so re-opening a quotation saved before the image store
    existed still renders its photos. Returns None when the field is empty or the referenced
    file has gone missing, which callers already treat as "no image on this row".
    """
    return image_store.resolve_bytes(value)


def item_image(item):
    """The image field of a line item, preferring the ref and falling back to legacy inline
    data so old history records and freshly parsed items both work through one path."""
    return item.get('image_ref') or item.get('image_base64') or ""


def _prepare_thumbnail(img_bytes, box=THUMB_BOX, pad=True):
    """Normalizes an arbitrary product photo into a clean, undistorted thumbnail.

    `pad` letterboxes onto a white square, which Excel needs so every image cell is the same
    shape. Word sets pad=False: there the picture is placed in a flowing table cell, so the
    white bands just made rows taller and the product smaller for no benefit.

    Two real bugs this fixes: (1) forcing an image into a fixed width/height box without
    preserving aspect ratio stretches/squishes any non-square photo, and (2) CMYK or
    palette-mode source images (common from print-oriented exports) render with garish,
    inverted-looking colors if not normalized to RGB first. Both showed up as "weird
    images" in generated quotes.
    """
    pil_img = PILImage.open(io.BytesIO(img_bytes))

    if pil_img.mode == "RGBA":
        # Flatten transparency onto a white background rather than letting Excel show
        # whatever garbage color sits behind an unsupported alpha channel.
        flattened = PILImage.new("RGB", pil_img.size, (255, 255, 255))
        flattened.paste(pil_img, mask=pil_img.split()[-1])
        pil_img = flattened
    elif pil_img.mode != "RGB":
        pil_img = pil_img.convert("RGB")

    # Fit within the box preserving aspect ratio — no stretching, no cropping out part of
    # the product. LANCZOS because these end up printed.
    if pad:
        out_img = ImageOps.pad(pil_img, (box, box), method=PILImage.LANCZOS,
                               color=(255, 255, 255), centering=(0.5, 0.5))
    else:
        out_img = pil_img.copy()
        out_img.thumbnail((box, box), PILImage.LANCZOS)

    stream = io.BytesIO()
    # JPEG, not PNG: these are photographs, and PNG made every generated .xlsx/.docx several
    # times larger for no visible gain at print size. Quality 90 because this is the copy the
    # client actually receives.
    out_img.save(stream, format="JPEG", quality=90, optimize=True)
    stream.seek(0)
    return stream


def _load_logo_scaled(max_width=140, max_height=60):
    """Loads COMPANY['logo_path'] if configured and scales it to fit within a bounding box,
    preserving both aspect ratio and transparency (unlike product thumbnails, a logo should
    keep its native shape rather than get letterboxed onto a white square). Returns
    (stream, width, height) or None if no logo is configured / the file is missing.

    logo_path is normally a relative path (e.g. "assets/red_cube_logo.png") so it stays
    portable in company.json regardless of what directory the app is launched from — resolve
    it against this file's own directory rather than trusting the process's current working
    directory, which a desktop app can't guarantee.
    """
    logo_path = COMPANY.get("logo_path", "")
    if not logo_path:
        return None
    path_obj = Path(logo_path)
    if not path_obj.is_absolute():
        path_obj = Path(__file__).resolve().parent / path_obj
    if not path_obj.exists():
        return None
    try:
        pil_img = PILImage.open(str(path_obj))
        ratio = min(max_width / pil_img.width, max_height / pil_img.height, 1.0)
        w, h = max(1, int(pil_img.width * ratio)), max(1, int(pil_img.height * ratio))
        stream = io.BytesIO()
        pil_img.save(stream, format="PNG")
        stream.seek(0)
        return stream, w, h
    except Exception as e:
        print(f"Failed to load company logo '{logo_path}': {e}")
        return None


def _insert_excel_logo(ws, anchor_cell="F1"):
    """Places the configured logo in the top-right corner of the sheet, out of the way of
    the client/date block that templates typically put on the left."""
    result = _load_logo_scaled()
    if not result:
        return
    stream, w, h = result
    try:
        oxl_img = openpyxl.drawing.image.Image(stream)
        oxl_img.width, oxl_img.height = w, h
        ws.add_image(oxl_img, anchor_cell)
    except Exception as e:
        print(f"Failed to insert company logo into Excel: {e}")


# --- Excel generation -----------------------------------------------------------

def create_fallback_template(dest_path):
    """Creates the branding/header block (rows 1-9) from scratch when no template file exists."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Quotation"

    header_fill = PatternFill(start_color=COMPANY["primary_hex"], end_color=COMPANY["primary_hex"], fill_type="solid")
    font_title = Font(name="Georgia", size=18, bold=True, color=COMPANY["primary_hex"])
    font_subtitle = Font(name=DOC_FONT, size=10.5, italic=True, color="6B6B6B")
    # White on the near-black fill. This was accent red on near-black, which is barely legible
    # on screen and worse in print — the same pairing that was fixed in the Word header.
    font_header = Font(name=DOC_FONT, size=10, bold=True, color="FFFFFF")
    font_bold = Font(name=DOC_FONT, size=10, bold=True)
    font_normal = Font(name=DOC_FONT, size=10)
    align_center = Alignment(horizontal="center", vertical="center")

    thin_border = Border(
        left=Side(style='thin', color='D9D9D9'), right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'), bottom=Side(style='thin', color='D9D9D9')
    )

    accent_underline = PatternFill(start_color=COMPANY["accent_hex"], end_color=COMPANY["accent_hex"], fill_type="solid")
    for col in range(1, 8):
        ws.cell(row=1, column=col).fill = accent_underline
    ws.row_dimensions[1].height = 4

    ws['A2'] = COMPANY["name"]
    ws['A2'].font = font_title
    ws['A3'] = COMPANY.get("tagline") or "Quotation Sheet"
    ws['A3'].font = font_subtitle

    ws['A5'] = "Client Name:"
    ws['A5'].font = font_bold
    ws['D5'] = "Date:"          # D:E is merged at fill time so long labels are not clipped
    ws['D5'].font = font_bold
    ws['A6'] = "Venue:"
    ws['A6'].font = font_bold

    # Images get a dedicated column in the template itself rather than being appended later,
    # and the specification column is the widest thing on the sheet — the client reads the
    # spec to know what they are buying, so it is given the room the reference sheets give it.
    headers = ["S.No", "Item & Specifications", "Images", "Unit", "Qty",
               "Rate (AED)", "VAT 5%", "TOTAL (AED)"]
    for col_idx, text in enumerate(headers, start=1):
        cell = ws.cell(row=9, column=col_idx, value=text)
        cell.font = font_header
        cell.fill = header_fill
        cell.alignment = align_center
        cell.border = thin_border
    ws.row_dimensions[9].height = 30

    # Specs column gives width back to Images: the photo is what a client recognises the
    # product from, and the spec text wraps happily in a narrower column.
    for col, width in zip("ABCDEFGH", [6, 46, 30, 10, 8, 14, 12, 16]):
        ws.column_dimensions[col].width = width

    _insert_excel_logo(ws, anchor_cell="F1")

    wb.save(dest_path)


def _detect_template_columns(ws, max_scan_row=25):
    col_map = {}
    header_row = None
    for r in range(1, max_scan_row + 1):
        row_vals = [ws.cell(row=r, column=c).value for c in range(1, 15)]
        row_str = [str(v).strip().lower() if v is not None else "" for v in row_vals]
        if any("description" in v or "particulars" in v or "specification" in v for v in row_str):
            header_row = r
            for idx, val in enumerate(row_str, start=1):
                if "item #" in val or "item no" in val or "s.no" in val or "sr no" in val:
                    col_map["item_num"] = idx
                elif "description" in val or "particulars" in val or "specification" in val:
                    col_map["description"] = idx
                elif "image" in val or "photo" in val:
                    col_map["images"] = idx
                elif "vat" in val:
                    col_map["vat"] = idx
                elif "total" in val:
                    col_map["total"] = idx
                elif "qty" in val or "quantity" in val:
                    col_map["qty"] = idx
                elif "rate" in val or "unit price" in val:
                    col_map["rate"] = idx
                elif "unit" in val:
                    col_map["unit"] = idx
            break

    if "description" not in col_map:
        col_map = {"item_num": 1, "description": 2, "unit": 3, "qty": 4, "rate": 5, "vat": 6, "total": 7}
        header_row = header_row or 9

    return header_row, col_map


def _row_height_for(has_image, spec_lines, name_lines=1):
    """Row height in points, clearing whichever is taller: the photo or the text beside it.

    Excel will not auto-fit a wrapped cell that also holds a floating image, so it is computed.
    Photo height: px -> pt is 0.75, plus padding.

    name_lines matters as much as spec_lines: a long item name wraps in the specs column too,
    and counting only the spec lines clipped the second line of names like "Additional
    Transportation + Installation + Crew Cost" clean off the row.
    """
    image_h = (THUMB_BOX * 0.75 + 8) if has_image else 0
    text_h = 10 + (max(1, name_lines) * 14) + (spec_lines * 12)
    return max(image_h, text_h, 20)


def _style_row(ws, row, col_map, thin_border, has_image=False, spec_lines=0, band=False,
               name_lines=1):
    align_left_wrap = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_center = Alignment(horizontal="center", vertical="center")
    align_right = Alignment(horizontal="right", vertical="center")
    font_normal = Font(name=DOC_FONT, size=10)
    band_fill = PatternFill(start_color=BAND_FILL, end_color=BAND_FILL, fill_type="solid")

    for key, col in col_map.items():
        cell = ws.cell(row=row, column=col)
        # The description cell carries rich text with its own fonts; assigning a plain Font
        # here would flatten the bold name back into the body weight.
        if key != "description":
            cell.font = font_normal
        cell.border = thin_border
        if band:
            cell.fill = band_fill
        if key == "description":
            cell.alignment = align_left_wrap
        elif key in ("rate", "vat", "total"):
            cell.alignment = align_right
            cell.number_format = '#,##0.00'
        else:
            cell.alignment = align_center

    # The real Red Cube template (sample_quotes/Red Cube - Quotation Format.xlsx) uses a flat
    # 20pt row height throughout — forcing every row to 45pt regardless of content made every
    # generated quote look visibly bloated/off compared to the business's actual documents.
    ws.row_dimensions[row].height = _row_height_for(has_image, spec_lines, name_lines)


def _apply_print_setup(ws, last_row, last_col, header_row):
    """Forces the sheet onto a single page width when printed or exported to PDF.

    Without this the generated workbook paginates on Excel's defaults, which split the table
    vertically down the middle — the Qty/Rate/VAT/Total columns landed on a separate sheet of
    the PDF, so the first page a client opened showed the line items with no prices on them.
    """
    try:
        ws.page_setup.orientation = 'landscape'
        ws.page_setup.paperSize = ws.PAPERSIZE_A4
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0  # allow as many pages tall as needed, never wider

        ws.print_area = f"A1:{get_column_letter(last_col)}{last_row}"
        # Repeat the column headings on every page of a long quotation.
        ws.print_title_rows = f"{header_row}:{header_row}"

        ws.page_margins.left = ws.page_margins.right = 0.4
        ws.page_margins.top = ws.page_margins.bottom = 0.5
    except Exception as e:
        print(f"Could not apply print setup: {e}")


def generate_excel_dynamic(items, meta, template_path, output_path):
    """Populates a quotation Excel sheet with exactly len(items) rows plus a discount/summary
    block — never leaves stale empty rows from a fixed-size template."""
    wb = openpyxl.load_workbook(str(template_path), data_only=False)
    ws = wb.active

    header_row, col_map = _detect_template_columns(ws)

    thin_border = Border(
        left=Side(style='thin', color='D9D9D9'), right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'), bottom=Side(style='thin', color='D9D9D9')
    )
    font_bold = Font(name=DOC_FONT, size=10, bold=True)
    font_grand = Font(name=DOC_FONT, size=11.5, bold=True, color=COMPANY["primary_hex"])
    align_right = Alignment(horizontal="right", vertical="center")

    # Clear any pre-existing stale rows below the header from an old fixed template.
    if ws.max_row > header_row:
        stale_last_row = ws.max_row
        ws.delete_rows(header_row + 1, stale_last_row - header_row)
        # delete_rows shifts cell VALUES up but leaves the RowDimension height entries behind,
        # so the old template's 55pt rows would otherwise stick to whatever we write next and
        # balloon the summary/terms/signature block onto extra pages. Drop those height entries
        # so any row we don't explicitly size falls back to Excel's compact default.
        for stale_row in range(header_row + 1, stale_last_row + 1):
            if stale_row in ws.row_dimensions:
                del ws.row_dimensions[stale_row]

    # Overwrite the title/subtitle unconditionally rather than trusting whatever static text
    # is baked into the template file — otherwise company.json can say the right name and the
    # actual generated document still shows whatever text someone typed into the template once.
    ws.cell(row=2, column=1, value=COMPANY["name"]).font = Font(name="Georgia", size=16, bold=True, color=COMPANY["primary_hex"])
    ws.cell(row=3, column=1, value=COMPANY.get("tagline") or "Quotation Sheet").font = Font(name=DOC_FONT, size=10.5, italic=True, color="6B6B6B")

    quote_date = meta.get("quote_date", datetime.now().strftime("%Y-%m-%d"))
    valid_until = meta.get("valid_until") or compute_valid_until(quote_date)
    font_bold_header = Font(name=DOC_FONT, size=10, bold=True)

    # The right-hand meta labels ("Valid Until:", "Quote Ref:") are longer than the narrow Qty
    # column they sit above, and Excel clips rather than overflows once the value cell beside
    # them is filled. Rows 5-7 are above the table, so the label spans D:E harmlessly.
    for meta_row in (5, 6, 7):
        try:
            ws.merge_cells(start_row=meta_row, start_column=4, end_row=meta_row, end_column=5)
        except Exception:
            pass

    if ws.cell(row=5, column=1).value and "client" in str(ws.cell(row=5, column=1).value).lower():
        ws.cell(row=5, column=2, value=meta.get("client_name", ""))
    ws.cell(row=5, column=4, value="Date:").font = font_bold_header
    ws.cell(row=5, column=6, value=quote_date)

    # Write the Venue/Valid Until label + value unconditionally rather than only when the
    # template already has that label — the production template on file turned out to have
    # no "Venue:" field at all, so venue was silently missing from every generated quote.
    ws.cell(row=6, column=1, value="Venue:").font = font_bold_header
    ws.cell(row=6, column=2, value=meta.get("venue", ""))
    ws.cell(row=6, column=4, value="Valid Until:").font = font_bold_header
    ws.cell(row=6, column=6, value=valid_until)

    pm_name = meta.get("pm_name") or COMPANY.get("pm_name", "")
    if pm_name:
        ws.cell(row=7, column=1, value="Prepared By:").font = font_bold_header
        ws.cell(row=7, column=2, value=pm_name)

    quote_ref = meta.get("quote_ref")
    if quote_ref:
        ws.cell(row=7, column=4, value="Quote Ref:").font = font_bold_header
        ws.cell(row=7, column=6, value=quote_ref).font = Font(
            name=DOC_FONT, size=10, bold=True, color=COMPANY["accent_hex"])

    desc_col = col_map["description"]
    unit_col = col_map.get("unit")
    qty_col = col_map.get("qty")
    rate_col = col_map.get("rate")
    vat_col = col_map.get("vat")
    tot_col = col_map.get("total")
    img_col = col_map.get("images")

    # The real Red Cube template has no "Images" column at all — without this, any photo a
    # line item carries (from the photo library, a PDF/Excel extraction, or a web search) had
    # nowhere to go in the generated sheet and was silently dropped. Append a trailing Image
    # column rather than inserting one mid-sheet, so none of the existing column letters used
    # by the meta rows above (Date/Venue/Valid Until/Prepared By) shift and break.
    if not img_col and any(item_image(item) for item in items):
        img_col = max(col_map.values()) + 1
        img_col_letter = get_column_letter(img_col)
        header_cell = ws.cell(row=header_row, column=img_col, value="Image")
        ref_header_cell = ws.cell(row=header_row, column=desc_col)
        header_cell.font = ref_header_cell.font.copy()
        header_cell.fill = ref_header_cell.fill.copy()
        header_cell.alignment = Alignment(horizontal="center", vertical="center")
        header_cell.border = ref_header_cell.border.copy()
        # ~7px per width unit, so the column tracks the photo instead of cropping it.
        ws.column_dimensions[img_col_letter].width = THUMB_BOX / 7.0 + 2

    current_row = header_row + 1
    for idx, item in enumerate(items, 1):
        if "item_num" in col_map:
            ws.cell(row=current_row, column=col_map["item_num"], value=idx)

        # Name in bold, specification lines beneath it, in a single cell — openpyxl rich text
        # is the only way to vary weight inside one cell, and splitting them into two columns
        # would break the template's column map.
        name, spec_lines = split_name_and_spec(item.get('description', ''))
        desc_cell = ws.cell(row=current_row, column=desc_col)
        if spec_lines:
            desc_cell.value = CellRichText([
                TextBlock(InlineFont(rFont=DOC_FONT, sz=10.5, b=True), name),
                TextBlock(InlineFont(rFont=DOC_FONT, sz=9, color="595959"),
                          "\n" + "\n".join(spec_lines)),
            ])
        else:
            desc_cell.value = CellRichText([
                TextBlock(InlineFont(rFont=DOC_FONT, sz=10.5, b=True), name),
            ])
        if unit_col:
            ws.cell(row=current_row, column=unit_col, value=item.get('unit', 'Pcs'))
        if qty_col:
            ws.cell(row=current_row, column=qty_col, value=float(item.get('qty', 0) or 0))
        if rate_col:
            ws.cell(row=current_row, column=rate_col, value=float(item.get('rate', 0) or 0))

        if qty_col and rate_col:
            qty_letter = get_column_letter(qty_col)
            rate_letter = get_column_letter(rate_col)
            if vat_col:
                ws.cell(row=current_row, column=vat_col, value=f"=ROUND({qty_letter}{current_row}*{rate_letter}{current_row}*0.05,2)")
            if tot_col:
                # Matches the real template's own formula (qty*rate*1.05) — the per-row TOTAL
                # column is VAT-inclusive there. Writing plain qty*rate here under-totaled every
                # line item by its VAT amount versus what the business's actual quotes show.
                ws.cell(row=current_row, column=tot_col, value=f"=ROUND({qty_letter}{current_row}*{rate_letter}{current_row}*1.05,2)")

        # Row height has to be known before the image is placed, because an Excel picture is a
        # floating shape anchored to a cell corner — it does not inherit the cell's vertical
        # centering the way text does. Computing the height first lets the anchor be offset so
        # the photo sits centred against the specification text beside it.
        col_chars = max(20, int(ws.column_dimensions[get_column_letter(desc_col)].width or 40))
        wrapped = sum(max(1, (len(ln) // col_chars) + 1) for ln in spec_lines)
        # The name is bold and a shade larger, so it wraps sooner than the spec text does.
        name_wrapped = max(1, (len(name) // max(12, int(col_chars * 0.85))) + 1)
        has_image = bool(item_image(item) and img_col)
        row_height = _row_height_for(has_image, wrapped, name_wrapped)

        row_has_image = False
        if has_image:
            try:
                img_bytes = load_image_bytes(item_image(item))
                # Rendered at DOC_IMAGE_PX but displayed at THUMB_BOX: Excel keeps the full
                # resolution in the file, so zooming or printing stays sharp.
                tmp_stream = _prepare_thumbnail(img_bytes, DOC_IMAGE_PX)
                oxl_img = openpyxl.drawing.image.Image(tmp_stream)
                oxl_img.width = THUMB_BOX
                oxl_img.height = THUMB_BOX

                row_px = row_height / 0.75                       # pt -> px
                col_px = (ws.column_dimensions[get_column_letter(img_col)].width or 20) * 7.0
                y_off = max(0, (row_px - THUMB_BOX) / 2.0)
                x_off = max(0, (col_px - THUMB_BOX) / 2.0)
                marker = AnchorMarker(col=img_col - 1, colOff=pixels_to_EMU(x_off),
                                      row=current_row - 1, rowOff=pixels_to_EMU(y_off))
                oxl_img.anchor = OneCellAnchor(
                    _from=marker,
                    ext=XDRPositiveSize2D(pixels_to_EMU(THUMB_BOX), pixels_to_EMU(THUMB_BOX)),
                )
                ws.add_image(oxl_img)
                row_has_image = True
            except Exception as e:
                IMAGE_FAILURES.append(f"row {idx}: {e}")
                print(f"Failed to insert image into Excel: {e}")

        _style_row(ws, current_row, col_map, thin_border, has_image=row_has_image,
                   spec_lines=wrapped, band=(idx % 2 == 0), name_lines=name_wrapped)
        current_row += 1

    totals = compute_totals(items, meta.get("discount_type"), meta.get("discount_value", 0))
    label_col = desc_col
    value_col = tot_col or (rate_col + 1 if rate_col else desc_col + 1)

    def write_summary_row(label, value, grand=False, fill_hex=None):
        """One banded totals row spanning the money columns, matching the reference sheets
        where each total is a differently-coloured band rather than loose text."""
        nonlocal current_row
        if grand:
            font = Font(name=DOC_FONT, size=12, bold=True, color="FFFFFF")
        else:
            font = Font(name=DOC_FONT, size=10, bold=True, color=COMPANY["primary_hex"])
        fill = PatternFill(start_color=fill_hex, end_color=fill_hex, fill_type="solid") if fill_hex else None

        cell_label = ws.cell(row=current_row, column=label_col, value=label)
        cell_label.font = font
        cell_label.alignment = align_right
        cell_value = ws.cell(row=current_row, column=value_col, value=value)
        cell_value.font = font
        cell_value.alignment = align_right
        cell_value.number_format = '#,##0.00'
        # Fill the whole span so the band reads as one bar, not two disconnected cells.
        for col in range(label_col, value_col + 1):
            c = ws.cell(row=current_row, column=col)
            if fill:
                c.fill = fill
            c.border = thin_border
        ws.row_dimensions[current_row].height = 22 if grand else 18
        current_row += 1

    current_row += 1
    write_summary_row("Subtotal (AED)", totals["subtotal"], fill_hex=SUBTOTAL_FILL)
    if totals["discount_amount"] > 0:
        write_summary_row("Discount (AED)", -totals["discount_amount"], fill_hex=DISCOUNT_FILL)
    write_summary_row("VAT 5% (AED)", totals["vat"], fill_hex=SUBTOTAL_FILL)
    write_summary_row("Grand Total (AED)", totals["grand_total"], grand=True,
                      fill_hex=COMPANY["accent_hex"])

    # --- Terms & payment schedule (configurable via terms.json) ---
    current_row += 2
    ws.cell(row=current_row, column=label_col, value="Terms & Payment Schedule").font = Font(name=DOC_FONT, size=10.5, bold=True, color=COMPANY["primary_hex"])
    current_row += 1
    ws.cell(row=current_row, column=label_col, value=f"This quotation is valid until {valid_until}.").font = Font(name=DOC_FONT, size=9, italic=True)
    current_row += 1
    for term in TERMS_CONFIG.get("payment_terms", []):
        ws.cell(row=current_row, column=label_col, value=f"• {term}").font = Font(name=DOC_FONT, size=9)
        current_row += 1

    # --- Client acceptance / signature block ---
    current_row += 2
    ws.cell(row=current_row, column=label_col, value="Client Acceptance").font = Font(
        name=DOC_FONT, size=10.5, bold=True, color=COMPANY["primary_hex"])
    current_row += 1
    ws.cell(
        row=current_row, column=label_col,
        value="By signing below, the client accepts the scope, quantities and pricing above and authorises works to proceed.",
    ).font = Font(name=DOC_FONT, size=9, italic=True)
    current_row += 2

    # Labels were dropped into the very wide specs column while their rules landed in whatever
    # column happened to come next, so the two pairs never lined up — the "wiggly" look. Each
    # label is now right-aligned hard against its rule, and each rule is a merged span of a
    # fixed width, so both rows are identical regardless of the template's column widths.
    sign_border = Border(bottom=Side(style='thin', color='808080'))
    last_col = max(col_map.values())
    label_l, rule_l = 2, (3, 4)          # B label | C:D rule
    label_r, rule_r = (5, 6), (7, last_col)  # E:F label | G:last rule

    for left_label, right_label in (("Client Name:", "Authorised Signature:"), ("Date:", "Company Stamp:")):
        for span in (rule_l, label_r, rule_r):
            try:
                ws.merge_cells(start_row=current_row, start_column=span[0],
                               end_row=current_row, end_column=span[1])
            except Exception:
                pass

        cl = ws.cell(row=current_row, column=label_l, value=left_label)
        cl.font = font_bold
        cl.alignment = Alignment(horizontal="right", vertical="center")
        cr = ws.cell(row=current_row, column=label_r[0], value=right_label)
        cr.font = font_bold
        cr.alignment = Alignment(horizontal="right", vertical="center")
        # Border must be set on every cell of a merged span, not just the anchor, or the rule
        # only draws under the first column.
        for start, end in (rule_l, rule_r):
            for c in range(start, end + 1):
                ws.cell(row=current_row, column=c).border = sign_border
        ws.row_dimensions[current_row].height = 20
        current_row += 2

    _insert_excel_logo(ws, anchor_cell="F1")

    # Column A carries both the Item # and the meta labels ("Client Name:", "Prepared By:").
    # At the template's default width those labels are clipped as soon as the value cell
    # beside them is filled, so it gets just enough room to show them in full.
    if ws.column_dimensions['A'].width and ws.column_dimensions['A'].width < 13:
        ws.column_dimensions['A'].width = 13

    _apply_print_setup(ws, last_row=current_row, last_col=max(col_map.values()), header_row=header_row)

    wb.save(str(output_path))
    return totals


# --- Word generation ---------------------------------------------------------------

def _set_cell_shading(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


# --- Word layout helpers -----------------------------------------------------------
# The generated .docx is the artefact a client actually receives, so it is laid out
# deliberately rather than left to python-docx defaults: explicit column widths, borderless
# meta blocks, banded rows, and a single accent colour used only for emphasis.

MUTED_HEX = "6E6A66"
RULE_HEX = "DCD8D0"
# Band colour lives in the shared palette near the top so Word and Excel cannot drift apart.


def _set_cell_border(cell, **edges):
    """Applies individual cell borders. python-docx exposes no API for this."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        spec = edges.get(edge)
        if not spec:
            continue
        tag = 'w:{}'.format(edge)
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        for key, val in spec.items():
            el.set(qn('w:{}'.format(key)), str(val))


def _no_borders(table):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, **{e: {"val": "nil"} for e in ("top", "left", "bottom", "right")})


def _set_col_widths(table, widths_cm):
    """Pins column widths.

    Two traps here: widths only stick when written onto *every cell* (not the column), and
    they only hold if the table layout is fixed — otherwise Word re-autofits and the widths
    are advisory. Call this AFTER every row has been added; rows appended later inherit
    nothing, which silently pushed the rightmost columns off the page.
    """
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def _write_cell(cell, text, *, bold=False, italic=False, size=9.5, color=None,
                align=WD_ALIGN_PARAGRAPH.LEFT, font="Calibri", caps=False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(str(text).upper() if caps else str(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color is not None:
        run.font.color.rgb = color
    return para


def _rule(doc, hex_color, size=8, space_after=6):
    """A horizontal rule drawn as a paragraph bottom border — the reliable way in Word."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(space_after)
    pPr = para._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color)
    borders.append(bottom)
    pPr.append(borders)
    return para


def _section_heading(doc, text, accent_rgb):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = accent_rgb
    _set_char_spacing(run, 60)
    _rule(doc, RULE_HEX, size=6, space_after=6)
    return para


def _set_char_spacing(run, twentieths):
    """Letter-spacing for small-caps style headings; no python-docx API for this."""
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), str(twentieths))
    run._r.get_or_add_rPr().append(spacing)


def _add_page_number(paragraph):
    for kind, text in (("begin", None), (None, "PAGE"), ("end", None)):
        run = paragraph.add_run()
        if kind:
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), kind)
            run._r.append(fld)
        else:
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = text
            run._r.append(instr)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(*bytes.fromhex(MUTED_HEX))


def _build_footer(section, left_text):
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(left_text + "    |    Page ")
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(*bytes.fromhex(MUTED_HEX))
    _add_page_number(para)


def generate_word_dynamic(items, meta, output_path):
    """Generates the client-facing Word quotation.

    Laid out as a document someone is expected to sign, not a data dump: a branded masthead,
    a borderless detail block, a banded line-item table with the photo large enough to
    actually identify the product, an emphasised total, then terms and an acceptance block.
    """
    doc = Document()
    primary_rgb = RGBColor(*bytes.fromhex(COMPANY["primary_hex"]))
    accent_rgb = RGBColor(*bytes.fromhex(COMPANY["accent_hex"]))
    muted_rgb = RGBColor(*bytes.fromhex(MUTED_HEX))
    white_rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # A4 explicitly: python-docx defaults to US Letter, which is the wrong paper for a UAE
    # business and shifts every column width assumption below. Margins are tighter than
    # Word's 1-inch default so the table has room to breathe at this column count.
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)

    quote_date = meta.get("quote_date", datetime.now().strftime("%Y-%m-%d"))
    valid_until = meta.get("valid_until") or compute_valid_until(quote_date)
    pm_name = meta.get("pm_name") or COMPANY.get("pm_name", "")
    quote_ref = meta.get("quote_ref", "")
    # The same generator produces quotations and invoices. They differ in wording, not in
    # layout, so the labels are driven from meta rather than duplicating 200 lines.
    doc_kind = str(meta.get("doc_kind") or "QUOTATION").upper()
    is_invoice = doc_kind == "INVOICE"

    # --- Masthead: company identity left, logo right ---------------------------------
    head = doc.add_table(rows=1, cols=2)
    _no_borders(head)
    _set_col_widths(head, [12.6, 5.0])
    head.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    head.rows[0].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    logo = _load_logo_scaled(max_width=300, max_height=140)
    logo_cell = head.rows[0].cells[1]
    logo_cell.text = ""
    if logo:
        logo_stream, logo_w, _ = logo
        # Size the logo from its own pixel width so it is never upscaled into a blurry block:
        # a small crisp mark reads as more professional than a large soft one. Drop a
        # higher-resolution file into assets/ and it will render larger automatically.
        target_cm = max(1.7, min(3.2, (logo_w / LOGO_TARGET_DPI) * 2.54))
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo_para.add_run().add_picture(logo_stream, width=Cm(target_cm))

    ident = head.rows[0].cells[0]
    ident.text = ""
    name_para = ident.paragraphs[0]
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(COMPANY["name"])
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.name = "Georgia"
    name_run.font.color.rgb = primary_rgb
    _set_char_spacing(name_run, 30)

    tag_para = ident.add_paragraph()
    tag_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tag_para.paragraph_format.space_before = Pt(1)
    tag_run = tag_para.add_run(
        "Tax Invoice" if is_invoice
        else (COMPANY.get("tagline") or "Quotation & Cost Estimate"))
    tag_run.font.size = Pt(9)
    tag_run.italic = True
    tag_run.font.color.rgb = muted_rgb

    for line in (COMPANY.get("address"), COMPANY.get("contact")):
        if line:
            c = ident.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.LEFT
            c.paragraph_format.space_before = Pt(0)
            r = c.add_run(str(line))
            r.font.size = Pt(8)
            r.font.color.rgb = muted_rgb

    _rule(doc, COMPANY["accent_hex"], size=18, space_after=10)

    # --- Detail block: two label/value pairs per row, no visible grid -----------------
    if is_invoice:
        details = [("Invoice No", quote_ref), ("Date", quote_date),
                   ("Client", meta.get("client_name", "")), ("Venue", meta.get("venue", "")),
                   ("Payment Due", meta.get("due_date", "")), ("Prepared By", pm_name)]
    else:
        details = [("Quote Ref", quote_ref), ("Date", quote_date),
                   ("Client", meta.get("client_name", "")), ("Venue", meta.get("venue", "")),
                   ("Valid Until", valid_until), ("Prepared By", pm_name)]
    details = [(k, v) for k, v in details if str(v).strip()]

    rows = (len(details) + 1) // 2
    meta_table = doc.add_table(rows=rows, cols=4)
    _no_borders(meta_table)
    _set_col_widths(meta_table, [2.6, 6.2, 2.6, 6.2])
    for i, (label, value) in enumerate(details):
        r, c = i // 2, (i % 2) * 2
        _write_cell(meta_table.rows[r].cells[c], label, size=8, bold=True,
                    color=muted_rgb, caps=True)
        _write_cell(meta_table.rows[r].cells[c + 1], value, size=10, bold=True,
                    color=primary_rgb)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # --- Line items -------------------------------------------------------------------
    headers = ["S.No", "Item & Specifications", "Image", "Unit", "Qty", "Rate (AED)", "Total (AED)"]
    # Sums to 17.6cm against 17.8cm of usable A4 width — leaving no slack pushed the money
    # columns off the right edge of the page. The spec column is the widest: it is what the
    # client actually reads to know what they are buying.
    # Width moved from specs to the image, matching the Excel sheet. Still sums to 17.6cm.
    widths = [1.2, 4.6, 5.0, 1.3, 0.9, 2.3, 2.3]
    table = doc.add_table(rows=1, cols=len(headers))

    # Mirrors the header order: S.No, Item & Specifications, Image, Unit, Qty, Rate, Total.
    aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
              WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
              WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT]
    SPEC_COL, IMAGE_COL = 1, 2

    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        # White on the brand near-black. The previous accent-red-on-near-black header was
        # close to unreadable in print.
        _write_cell(cell, text, bold=True, size=8.5, color=white_rgb, align=aligns[i], caps=True)
        _set_cell_shading(cell, COMPANY["primary_hex"])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(8.5)

    for idx, item in enumerate(items, 1):
        cells = table.add_row().cells
        qty = float(item.get('qty', 0) or 0)
        rate = float(item.get('rate', 0) or 0)
        # VAT-inclusive per-row total, matching the real Excel template's own convention
        # (qty*rate*1.05) so Word and Excel outputs agree with each other and with how the
        # business's actual quotes have always been calculated.
        line_total = round(qty * rate * 1.05, 2)

        name, spec_lines = split_name_and_spec(item.get('description', ''))

        values = [str(idx), None, None, str(item.get('unit', 'Pcs')),
                  "{:g}".format(qty), "{:,.2f}".format(rate), "{:,.2f}".format(line_total)]
        for i, val in enumerate(values):
            if val is None:
                continue
            _write_cell(cells[i], val, size=9.5, align=aligns[i],
                        bold=(i == 6), color=primary_rgb if i == 6 else None)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Name in bold, then the specification beneath it. The spec is what tells the client
        # what they are actually buying, so it is kept in full — the original code dropped
        # everything after the first newline.
        spec_cell = cells[SPEC_COL]
        spec_cell.text = ""
        head_para = spec_cell.paragraphs[0]
        head_para.paragraph_format.space_before = Pt(4)
        head_para.paragraph_format.space_after = Pt(2)
        head_run = head_para.add_run(name or "—")
        head_run.bold = True
        head_run.font.size = Pt(10)
        head_run.font.name = DOC_FONT
        head_run.font.color.rgb = primary_rgb
        for extra in spec_lines:
            p = spec_cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(extra)
            r.font.size = Pt(8.5)
            r.font.name = DOC_FONT
            r.font.color.rgb = muted_rgb
        spec_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if item_image(item):
            try:
                img_bytes = load_image_bytes(item_image(item))
                thumb_stream = _prepare_thumbnail(img_bytes, DOC_IMAGE_PX, pad=False)
                cells[IMAGE_COL].text = ""
                pic_para = cells[IMAGE_COL].paragraphs[0]
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_para.paragraph_format.space_before = Pt(4)
                pic_para.paragraph_format.space_after = Pt(4)
                pic_para.add_run().add_picture(thumb_stream, width=Cm(WORD_IMAGE_CM))
            except Exception as e:
                IMAGE_FAILURES.append("row {}: {}".format(idx, e))
                print("Failed to insert image into Word doc: {}".format(e))
        cells[IMAGE_COL].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Banded rows plus a hairline separator: readable across a wide table without the
        # heavy full grid the default "Table Grid" style draws.
        for cell in cells:
            if idx % 2 == 0:
                _set_cell_shading(cell, BAND_FILL)
            _set_cell_border(cell, bottom={"val": "single", "sz": 4, "color": RULE_HEX})

    # Applied only now that every row exists — see _set_col_widths.
    _set_col_widths(table, widths)

    # --- Totals: right-aligned block, grand total carries the accent -------------------
    totals = compute_totals(items, meta.get("discount_type"), meta.get("discount_value", 0))

    # Same colour language as the Excel sheet: each total is its own band.
    summary_rows = [("Subtotal (AED)", totals["subtotal"], False, SUBTOTAL_FILL)]
    if totals["discount_amount"] > 0:
        summary_rows.append(("Discount (AED)", -totals["discount_amount"], False, DISCOUNT_FILL))
    summary_rows.append(("VAT 5% (AED)", totals["vat"], False, SUBTOTAL_FILL))
    summary_rows.append(("Grand Total (AED)", totals["grand_total"], True, COMPANY["accent_hex"]))
    if is_invoice and float(meta.get("amount_paid") or 0) > 0:
        # Only shown once something has been paid. On an untouched invoice a "Balance Due"
        # identical to the grand total is noise, and a zero "Amount Paid" reads as an error.
        summary_rows.append(("Amount Paid (AED)", float(meta["amount_paid"]), False, SUBTOTAL_FILL))
        summary_rows.append(("Balance Due (AED)", float(meta.get("balance_due") or 0),
                             True, COMPANY["accent_hex"]))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)

    sum_table = doc.add_table(rows=len(summary_rows), cols=2)
    sum_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _no_borders(sum_table)
    _set_col_widths(sum_table, [5.4, 3.4])
    for i, (label, value, is_total, fill_hex) in enumerate(summary_rows):
        lc, vc = sum_table.rows[i].cells[0], sum_table.rows[i].cells[1]
        _write_cell(lc, label, bold=True, size=11 if is_total else 9.5,
                    color=white_rgb if is_total else primary_rgb,
                    align=WD_ALIGN_PARAGRAPH.RIGHT, caps=is_total)
        _write_cell(vc, "{:,.2f}".format(value), bold=True, size=13 if is_total else 9.5,
                    color=white_rgb if is_total else primary_rgb,
                    align=WD_ALIGN_PARAGRAPH.RIGHT)
        for cell in (lc, vc):
            _set_cell_shading(cell, fill_hex)

    # --- Terms & payment schedule (configurable via terms.json) -----------------------
    _section_heading(doc, "Terms & Payment Schedule", accent_rgb)

    validity = doc.add_paragraph()
    validity.paragraph_format.space_after = Pt(4)
    vr = validity.add_run("This quotation is valid until {}.".format(valid_until))
    vr.italic = True
    vr.font.size = Pt(9)
    vr.font.color.rgb = muted_rgb

    for term in TERMS_CONFIG.get("payment_terms", []):
        p = doc.add_paragraph(str(term), style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(9)

    # --- Client acceptance ------------------------------------------------------------
    _section_heading(doc, "Client Acceptance", accent_rgb)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    ir = intro.add_run(
        "By signing below, the client confirms acceptance of the scope, quantities and "
        "pricing set out in this quotation, and authorises works to proceed."
    )
    ir.font.size = Pt(9)
    ir.font.color.rgb = muted_rgb

    sign_table = doc.add_table(rows=2, cols=2)
    _no_borders(sign_table)
    _set_col_widths(sign_table, [8.8, 8.8])
    sign_labels = [("Client Name", "Authorised Signature"), ("Date", "Company Stamp")]
    for r, pair in enumerate(sign_labels):
        for c, label in enumerate(pair):
            cell = sign_table.rows[r].cells[c]
            _write_cell(cell, label, bold=True, size=8, color=muted_rgb, caps=True)
            # Empty ruled line for the handwritten entry. Kept tight: an extra few points here
            # spilled the block onto a second, otherwise-blank page.
            line = cell.add_paragraph()
            line.paragraph_format.space_before = Pt(10)
            line.paragraph_format.space_after = Pt(2)
            _set_cell_border(cell, bottom={"val": "single", "sz": 4, "color": MUTED_HEX})

    footer_bits = [COMPANY["name"]]
    if quote_ref:
        footer_bits.append("{} {}".format("Invoice" if is_invoice else "Quotation", quote_ref))
    _build_footer(doc.sections[0], "    |    ".join(footer_bits))

    doc.save(str(output_path))
    return totals
